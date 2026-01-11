"""
SASTRE Document Export

Export investigation documents to various formats.
Supports Word (docx) and PDF with proper formatting.
"""

import re
from dataclasses import dataclass
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime

from .sections import Section, SectionState


@dataclass
class ExportConfig:
    """Configuration for document export."""
    include_metadata: bool = False         # Include machine-readable metadata
    include_footnotes: bool = True
    include_surprising_ands: bool = True
    format_entity_profiles: bool = True    # Core/Shell/Halo formatting
    style_template: Optional[str] = None   # Path to Word template

    # Header/footer
    header_text: Optional[str] = None
    footer_text: Optional[str] = None
    page_numbers: bool = True

    # Confidentiality
    confidential_marking: Optional[str] = None  # "CONFIDENTIAL", "PRIVILEGED", etc.


@dataclass
class ExportResult:
    """Result of document export."""
    success: bool
    output_path: Optional[Path]
    format: str
    page_count: int = 0
    word_count: int = 0
    error: Optional[str] = None


class MarkdownRenderer:
    """Render document to clean markdown."""

    @staticmethod
    def render(
        title: str,
        sections: List[Section],
        footnotes: List[str],
        config: ExportConfig
    ) -> str:
        """Render document as markdown."""
        lines = []

        # Title
        lines.append(f"# {title}")
        lines.append("")

        # Metadata (optional)
        if config.include_metadata:
            lines.append("<!-- DOCUMENT_META")
            lines.append(f"exported: {datetime.now().isoformat()}")
            lines.append("-->")
            lines.append("")

        # Confidential marking
        if config.confidential_marking:
            lines.append(f"**{config.confidential_marking}**")
            lines.append("")

        # Sections
        for section in sections:
            # Skip empty sections unless they have gaps
            if section.state == SectionState.EMPTY and not section.gaps:
                continue

            lines.append(section.header)
            lines.append("")

            if section.content:
                lines.append(section.content)
                lines.append("")

        # Footnotes
        if config.include_footnotes and footnotes:
            lines.append("---")
            lines.append("")
            lines.append("## References")
            lines.append("")
            for fn in footnotes:
                lines.append(fn)

        return '\n'.join(lines)


class WordExporter:
    """Export document to Word (docx) format."""

    def __init__(self, config: Optional[ExportConfig] = None):
        self.config = config or ExportConfig()

    def export(
        self,
        title: str,
        sections: List[Section],
        footnotes: List[str],
        output_path: Path
    ) -> ExportResult:
        """
        Export document to Word format.

        Requires python-docx library.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return ExportResult(
                success=False,
                output_path=None,
                format="docx",
                error="python-docx library not installed. Install with: pip install python-docx"
            )

        try:
            # Create document
            if self.config.style_template:
                doc = Document(self.config.style_template)
            else:
                doc = Document()

            # Add confidential marking if configured
            if self.config.confidential_marking:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(self.config.confidential_marking)
                run.bold = True
                run.font.size = Pt(14)

            # Title
            doc.add_heading(title, level=0)

            # Sections
            word_count = 0
            for section in sections:
                if section.state == SectionState.EMPTY and not section.gaps:
                    continue

                # Section header (H2)
                doc.add_heading(section.clean_header, level=1)

                # Section content
                if section.content:
                    self._add_content(doc, section.content)
                    word_count += len(section.content.split())

            # Footnotes
            if self.config.include_footnotes and footnotes:
                doc.add_page_break()
                doc.add_heading("References", level=1)

                for fn in footnotes:
                    # Parse footnote: [^N]: text
                    match = re.match(r'\[\^(\d+)\]:\s*(.+)', fn)
                    if match:
                        p = doc.add_paragraph()
                        p.add_run(f"[{match.group(1)}] ").bold = True
                        p.add_run(match.group(2))

            # Save
            doc.save(str(output_path))

            return ExportResult(
                success=True,
                output_path=output_path,
                format="docx",
                word_count=word_count,
            )

        except Exception as e:
            return ExportResult(
                success=False,
                output_path=None,
                format="docx",
                error=str(e)
            )

    def _add_content(self, doc, content: str) -> None:
        """Add markdown content to Word document."""
        from docx.shared import Pt

        lines = content.split('\n')
        current_para = None

        for line in lines:
            line = line.strip()

            if not line:
                current_para = None
                continue

            # H3 header
            if line.startswith('### '):
                doc.add_heading(line[4:], level=2)
                current_para = None

            # Bold section (**text**)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line[2:-2]).bold = True
                current_para = None

            # Bullet point
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')

                # Handle footnote references
                text = line[2:]
                if '[^' in text:
                    # Parse footnote: text [^N]
                    parts = re.split(r'\[\^(\d+)\]', text)
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            p.add_run(part)
                        else:
                            run = p.add_run(f"[{part}]")
                            run.font.superscript = True
                else:
                    p.add_run(text)

                current_para = None

            # Regular paragraph
            else:
                if current_para is None:
                    current_para = doc.add_paragraph()

                # Handle inline formatting
                if '[^' in line:
                    parts = re.split(r'\[\^(\d+)\]', line)
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            current_para.add_run(part)
                        else:
                            run = current_para.add_run(f"[{part}]")
                            run.font.superscript = True
                else:
                    current_para.add_run(line + ' ')


class PDFExporter:
    """Export document to PDF format."""

    def __init__(self, config: Optional[ExportConfig] = None):
        self.config = config or ExportConfig()

    def export(
        self,
        title: str,
        sections: List[Section],
        footnotes: List[str],
        output_path: Path
    ) -> ExportResult:
        """
        Export document to PDF format.

        Strategy: Generate HTML then convert to PDF using weasyprint or pdfkit.
        """
        try:
            from weasyprint import HTML, CSS
        except ImportError:
            try:
                import pdfkit
                return self._export_with_pdfkit(title, sections, footnotes, output_path)
            except ImportError:
                return ExportResult(
                    success=False,
                    output_path=None,
                    format="pdf",
                    error="No PDF library available. Install weasyprint or pdfkit."
                )

        try:
            # Generate HTML
            html_content = self._to_html(title, sections, footnotes)

            # CSS styling
            css = CSS(string=self._get_css())

            # Generate PDF
            HTML(string=html_content).write_pdf(str(output_path), stylesheets=[css])

            return ExportResult(
                success=True,
                output_path=output_path,
                format="pdf",
            )

        except Exception as e:
            return ExportResult(
                success=False,
                output_path=None,
                format="pdf",
                error=str(e)
            )

    def _to_html(
        self,
        title: str,
        sections: List[Section],
        footnotes: List[str]
    ) -> str:
        """Convert document to HTML."""
        import html

        lines = ['<!DOCTYPE html>', '<html>', '<head>',
                 '<meta charset="UTF-8">',
                 f'<title>{html.escape(title)}</title>',
                 '</head>', '<body>']

        # Confidential marking
        if self.config.confidential_marking:
            lines.append(f'<div class="confidential">{html.escape(self.config.confidential_marking)}</div>')

        # Title
        lines.append(f'<h1>{html.escape(title)}</h1>')

        # Sections
        for section in sections:
            if section.state == SectionState.EMPTY and not section.gaps:
                continue

            lines.append(f'<h2>{html.escape(section.clean_header)}</h2>')

            if section.content:
                # Convert markdown to HTML (simplified)
                content_html = self._markdown_to_html(section.content)
                lines.append(content_html)

        # Footnotes
        if self.config.include_footnotes and footnotes:
            lines.append('<hr>')
            lines.append('<h2>References</h2>')
            lines.append('<div class="footnotes">')
            for fn in footnotes:
                match = re.match(r'\[\^(\d+)\]:\s*(.+)', fn)
                if match:
                    lines.append(f'<p><sup>{match.group(1)}</sup> {html.escape(match.group(2))}</p>')
            lines.append('</div>')

        lines.extend(['</body>', '</html>'])
        return '\n'.join(lines)

    def _markdown_to_html(self, content: str) -> str:
        """Simple markdown to HTML conversion."""
        import html

        result = []
        lines = content.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                result.append('<br>')
            elif line.startswith('### '):
                result.append(f'<h3>{html.escape(line[4:])}</h3>')
            elif line.startswith('**') and line.endswith('**'):
                result.append(f'<p><strong>{html.escape(line[2:-2])}</strong></p>')
            elif line.startswith('- '):
                text = html.escape(line[2:])
                # Convert footnote refs
                text = re.sub(r'\[\^(\d+)\]', r'<sup>\1</sup>', text)
                result.append(f'<li>{text}</li>')
            else:
                text = html.escape(line)
                text = re.sub(r'\[\^(\d+)\]', r'<sup>\1</sup>', text)
                result.append(f'<p>{text}</p>')

        return '\n'.join(result)

    def _get_css(self) -> str:
        """Get CSS for PDF styling."""
        return """
        body {
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.5;
            margin: 1in;
        }
        h1 {
            font-size: 18pt;
            text-align: center;
            margin-bottom: 24pt;
        }
        h2 {
            font-size: 14pt;
            margin-top: 18pt;
            border-bottom: 1px solid #333;
        }
        h3 {
            font-size: 12pt;
            font-weight: bold;
            margin-top: 12pt;
        }
        .confidential {
            text-align: center;
            font-weight: bold;
            color: red;
            margin-bottom: 24pt;
        }
        .footnotes {
            font-size: 10pt;
            margin-top: 24pt;
        }
        li {
            margin-left: 24pt;
        }
        sup {
            font-size: 8pt;
        }
        """

    def _export_with_pdfkit(
        self,
        title: str,
        sections: List[Section],
        footnotes: List[str],
        output_path: Path
    ) -> ExportResult:
        """Export using pdfkit as fallback."""
        import pdfkit

        try:
            html_content = self._to_html(title, sections, footnotes)
            pdfkit.from_string(html_content, str(output_path))

            return ExportResult(
                success=True,
                output_path=output_path,
                format="pdf",
            )
        except Exception as e:
            return ExportResult(
                success=False,
                output_path=None,
                format="pdf",
                error=str(e)
            )


def export_document(
    title: str,
    sections: List[Section],
    footnotes: List[str],
    output_path: Path,
    format: str = "markdown",
    config: Optional[ExportConfig] = None
) -> ExportResult:
    """
    Export document to specified format.

    Args:
        title: Document title
        sections: List of document sections
        footnotes: List of footnotes
        output_path: Where to save the export
        format: "markdown", "docx", or "pdf"
        config: Export configuration

    Returns:
        ExportResult with success status and details
    """
    config = config or ExportConfig()

    if format == "markdown" or format == "md":
        content = MarkdownRenderer.render(title, sections, footnotes, config)
        try:
            output_path.write_text(content)
            return ExportResult(
                success=True,
                output_path=output_path,
                format="markdown",
                word_count=len(content.split()),
            )
        except Exception as e:
            return ExportResult(
                success=False,
                output_path=None,
                format="markdown",
                error=str(e)
            )

    elif format == "docx" or format == "word":
        exporter = WordExporter(config)
        return exporter.export(title, sections, footnotes, output_path)

    elif format == "pdf":
        exporter = PDFExporter(config)
        return exporter.export(title, sections, footnotes, output_path)

    else:
        return ExportResult(
            success=False,
            output_path=None,
            format=format,
            error=f"Unknown format: {format}. Supported: markdown, docx, pdf"
        )
